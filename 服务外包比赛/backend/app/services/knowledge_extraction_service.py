"""
知识抽取服务包装器
整合 triplet_extractor 到主项目中，提供统一的接口
"""
import os
import json
import logging
import tempfile
from typing import Dict, Any, List, Optional
from datetime import datetime
from app.services.triplet_extractor import QwenExtractor
from app.models.triplets import TripletExtractionResult
from pydantic import ValidationError

logger = logging.getLogger(__name__)


class KnowledgeExtractionService:
    """知识抽取服务包装器"""
    
    def __init__(self, mode: str = "multimodal"):
        """
        初始化知识抽取服务
        
        Args:
            mode: 处理模式 ("legacy" 或 "multimodal")
        """
        self.mode = mode
        self.extractor = QwenExtractor()
        logger.info(f"知识抽取服务初始化完成，模式：{mode}")
    
    def extract_from_file(self, file_path: str, source_reference: str = None, top_event: str = None) -> Dict[str, Any]:
        """
        从文件提取知识三元组
        
        Args:
            file_path: 文件路径
            source_reference: 来源引用（文件名或顶事件名称）
            top_event: 顶事件名称（用于命名输出文件和作为来源引用）
            
        Returns:
            包含提取结果的字典
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在：{file_path}")
        
        # 读取文件内容
        file_content = self._read_file_content(file_path)
        
        # 优先使用顶事件作为来源引用，否则使用文件名
        source_ref = top_event if top_event else (source_reference or os.path.basename(file_path))
        
        # 使用 LLM 提取三元组
        logger.info(f"开始从文件提取三元组：{source_ref}")
        json_result = self.extractor.extract(file_content, source_reference=source_ref)
        
        # 解析和验证结果
        try:
            parsed_data = json.loads(json_result)
            validated_data = TripletExtractionResult(**parsed_data)
            
            # 转换为字典格式
            triplets = [triplet.model_dump() for triplet in validated_data.triplets]
            
            logger.info(f"成功提取 {len(triplets)} 个三元组")
            
            return {
                "success": True,
                "triplets": triplets,
                "raw_json": json_result,
                "file_path": file_path,
                "source_reference": source_reference or os.path.basename(file_path),
                "top_event": top_event
            }
            
        except (json.JSONDecodeError, ValidationError) as e:
            logger.error(f"解析或验证失败：{str(e)}")
            return {
                "success": False,
                "error": str(e),
                "raw_response": json_result,
                "file_path": file_path
            }
    
    def _read_file_content(self, file_path: str) -> str:
        """
        读取文件内容（支持多种格式）
        
        Args:
            file_path: 文件路径
            
        Returns:
            文件文本内容
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            # 文本文件
            if ext in ['.txt', '.md', '.json', '.csv']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            
            # PDF 文件
            elif ext == '.pdf':
                import pdfplumber
                text = ""
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                return text
            
            # Word 文档
            elif ext in ['.docx', '.doc']:
                try:
                    import docx
                    doc = docx.Document(file_path)
                    text = ""
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    
                    # 添加表格内容
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                text += cell.text + "\t"
                            text += "\n"
                        text += "\n"
                    
                    return text
                except ImportError:
                    raise ImportError("请安装：pip install python-docx")
            
            else:
                # 默认尝试以文本方式读取
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
                    
        except Exception as e:
            logger.error(f"读取文件失败：{str(e)}")
            raise RuntimeError(f"文件读取失败 {file_path}: {str(e)}")
    
    def save_triplets_to_file(self, triplets: List[Dict], output_dir: str = "data/triplets", filename: str = None, top_event: str = None) -> str:
        """
        将三元组保存到 JSON 文件
        
        Args:
            triplets: 三元组列表
            output_dir: 输出目录（默认：data/triplets）
            filename: 文件名（可选）
            top_event: 顶事件名称（优先用于命名）
            
        Returns:
            保存的文件路径
        """
        os.makedirs(output_dir, exist_ok=True)
        
        if filename is None:
            if top_event:
                # 使用顶事件命名（与 fault_tree 格式一致）
                safe_top_event = "".join([c for c in top_event if c.isalnum() or c in '-_']).strip()
                filename = f"triplets_{safe_top_event}.json"
            else:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"triplets_{timestamp}.json"
        
        output_path = os.path.join(output_dir, filename)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump({"triplets": triplets}, f, ensure_ascii=False, indent=4)
        
        logger.info(f"三元组已保存至：{output_path}")
        return output_path
    
    def extract_and_save(self, file_path: str, output_dir: str = "data/output", top_event: str = None) -> Dict[str, Any]:
        """
        提取三元组并保存到文件
        
        Args:
            file_path: 输入文件路径
            output_dir: 输出目录
            top_event: 顶事件名称（用于命名输出文件）
            
        Returns:
            包含提取结果和输出路径的字典
        """
        result = self.extract_from_file(file_path, top_event=top_event)
        
        if result["success"]:
            output_filename = f"{os.path.splitext(os.path.basename(file_path))[0]}_triplets.json"
            output_path = self.save_triplets_to_file(
                result["triplets"], 
                output_dir=output_dir,
                filename=output_filename,
                top_event=top_event
            )
            result["output_path"] = output_path
        else:
            result["output_path"] = None
        
        return result
